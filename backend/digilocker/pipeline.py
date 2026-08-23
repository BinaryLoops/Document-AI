"""
digilocker/pipeline.py — Document processing pipeline.

Orchestrates the full workflow:
  Upload → Scan → OCR → Classification → Verification → Encryption → Storage → Retrieval

Each step is clearly separated so it can be tested / replaced independently.
"""

from __future__ import annotations

import base64
import io
import json
import logging
import mimetypes
import os
from datetime import datetime, timezone
from typing import Any, Dict, Optional, Tuple

from digilocker.database import DocumentDatabase
from digilocker.dedup import DuplicateDetector, compute_sha256, compute_perceptual_hash
from digilocker.encryption import AES256Encryptor
from digilocker.models import (
    AuditAction,
    AuditRecord,
    DocumentCategory,
    DocumentLifecycle,
    DocumentVersion,
    LockerDocument,
    VerificationStatus,
)
from digilocker.preview import PreviewGenerator
from digilocker.scanner import MalwareScanner, create_scanner
from digilocker.vault import FileVault

logger = logging.getLogger(__name__)

# Map from classifier output → DocumentCategory
_CLASSIFIER_MAP: Dict[str, DocumentCategory] = {
    "identity proof":         DocumentCategory.PASSPORT,
    "passport":               DocumentCategory.PASSPORT,
    "driving licence":        DocumentCategory.DRIVING_LICENCE,
    "driving license":        DocumentCategory.DRIVING_LICENCE,
    "birth certificate":      DocumentCategory.BIRTH_CERTIFICATE,
    "certificate":            DocumentCategory.EDUCATION_CERTIFICATE,
    "education certificate":  DocumentCategory.EDUCATION_CERTIFICATE,
    "income certificate":     DocumentCategory.INCOME_CERTIFICATE,
    "land record":            DocumentCategory.LAND_RECORD,
    "fir":                    DocumentCategory.FIR,
    "court document":         DocumentCategory.COURT_ORDER,
    "court order":            DocumentCategory.COURT_ORDER,
}

# MIME type mapping
SUPPORTED_MIME = {
    "application/pdf":  ".pdf",
    "image/jpeg":       ".jpg",
    "image/png":        ".png",
    "image/jpg":        ".jpg",
    "application/vnd.openxmlformats-officedocument.wordprocessingml.document": ".docx",
}

# Auto-verification confidence threshold
_AUTO_VERIFY_THRESHOLD = float(os.environ.get("AUTO_VERIFY_THRESHOLD", "0.75"))

# Category schema dir
_CATEGORY_DIR = os.path.join(os.path.dirname(__file__), "categories")


class DocumentPipeline:
    """
    Runs documents through the full Digital Locker ingest pipeline.

    Each method corresponds to one pipeline step and can be called
    independently for testing.
    """

    def __init__(
        self,
        db: DocumentDatabase,
        vault: FileVault,
        encryptor: AES256Encryptor,
        scanner: Optional[MalwareScanner] = None,
    ):
        self.db = db
        self.vault = vault
        self.encryptor = encryptor
        self.scanner = scanner or create_scanner()
        self.dedup = DuplicateDetector()
        self.preview_gen = PreviewGenerator()

    # ── Full pipeline ────────────────────────────────────────────────────

    async def ingest(
        self,
        file_bytes: bytes,
        filename: str,
        owner: str,
        department: str = "",
        case_type: str = "",
        serial_number: str = "",
        mime_type: str = "",
        user_id: str = "",
        ip_address: str = "",
    ) -> LockerDocument:
        """
        Run the full pipeline: Upload → Scan → OCR → Classification →
        Verification → Encryption → Storage.

        Returns:
            The persisted LockerDocument.

        Raises:
            ValueError: unsupported format or malware detected.
        """

        # ── Step 1: Upload validation ────────────────────────────────────
        if not mime_type:
            mime_type = self._detect_mime(filename)
        if mime_type not in SUPPORTED_MIME:
            raise ValueError(
                f"Unsupported file type: {mime_type}. "
                f"Supported: {', '.join(SUPPORTED_MIME.keys())}"
            )

        file_hash = compute_sha256(file_bytes)
        logger.info(
            "Pipeline START — file='%s' size=%d hash=%s…",
            filename, len(file_bytes), file_hash[:12],
        )

        # ── Step 2: Malware scan ─────────────────────────────────────────
        scan_result = self.scanner.scan(file_bytes, filename)
        if not scan_result.is_clean:
            logger.error(
                "MALWARE DETECTED in '%s': %s", filename, scan_result.detail
            )
            raise ValueError(f"Malware detected: {scan_result.detail}")
        logger.info("Step 2/7 — Malware scan PASSED (%s)", scan_result.engine)

        # ── Step 3: Duplicate detection ──────────────────────────────────
        is_dup, dup_id = await self.dedup.check(file_bytes, mime_type, self.db)
        if is_dup:
            logger.info("Duplicate detected → original document %s", dup_id)

        # ── Step 4: OCR & text extraction ────────────────────────────────
        ocr_text, ocr_confidence, page_count = self._extract_text(
            file_bytes, mime_type, filename
        )
        logger.info(
            "Step 4/7 — OCR extracted %d chars (confidence=%.2f, pages=%d)",
            len(ocr_text), ocr_confidence, page_count,
        )

        # ── Step 5: Classification ───────────────────────────────────────
        doc_category, class_confidence = self._classify(ocr_text)
        combined_confidence = (ocr_confidence + class_confidence) / 2
        logger.info(
            "Step 5/7 — Classified as %s (confidence=%.2f)",
            doc_category.value, class_confidence,
        )

        # ── Step 6: Verification ─────────────────────────────────────────
        if combined_confidence >= _AUTO_VERIFY_THRESHOLD:
            verification = VerificationStatus.VERIFIED
        else:
            verification = VerificationStatus.PENDING
        logger.info("Step 6/7 — Verification status: %s", verification.value)

        # ── Step 7: Encryption + storage ─────────────────────────────────
        encrypted_blob, wrapped_dek = self.encryptor.encrypt(file_bytes)
        doc_id = None  # will be set by LockerDocument default
        doc = LockerDocument(
            owner=owner,
            department=department,
            document_type=doc_category,
            case_type=case_type,
            serial_number=serial_number,
            verification_status=verification,
            confidence_score=round(combined_confidence, 4),
            file_hash=file_hash,
            original_filename=filename,
            file_size=len(file_bytes),
            mime_type=mime_type,
            page_count=page_count,
            ocr_text=ocr_text,
            is_duplicate=is_dup,
            duplicate_of=dup_id,
            lifecycle=DocumentLifecycle.ACTIVE,
        )

        # Store encrypted blob
        vault_ref = self.vault.store(doc.document_id, encrypted_blob, wrapped_dek)
        doc.encryption_ref = vault_ref

        # Generate preview + thumbnail
        preview_bytes, thumb_bytes = self.preview_gen.generate(
            file_bytes, mime_type, filename
        )
        if preview_bytes:
            doc.preview_ref = self.vault.store_preview(doc.document_id, preview_bytes)
        if thumb_bytes:
            doc.thumbnail_ref = self.vault.store_thumbnail(doc.document_id, thumb_bytes)

        # Generate QR code
        doc.qr_code = self._generate_qr(doc.document_id)

        # Extract metadata from category schema
        doc.extracted_metadata = self._extract_metadata(
            ocr_text, doc_category
        )
        # Store perceptual hash for images
        if mime_type.startswith("image/"):
            phash = compute_perceptual_hash(file_bytes)
            if phash:
                doc.extracted_metadata["perceptual_hash"] = phash

        # ── Persist to database ──────────────────────────────────────────
        await self.db.insert_document(doc)

        # Create initial version record
        version = DocumentVersion(
            document_id=doc.document_id,
            version=1,
            file_hash=file_hash,
            encryption_ref=vault_ref,
            file_size=len(file_bytes),
            change_summary="Initial upload",
            created_by=user_id or owner,
        )
        await self.db.insert_version(version)

        # Audit log (per-document trail)
        await self.db.log_audit(AuditRecord(
            document_id=doc.document_id,
            user_id=user_id or owner,
            action=AuditAction.UPLOAD,
            detail=f"Uploaded {filename} ({mime_type}, {len(file_bytes)} bytes)",
            ip_address=ip_address,
        ))

        # Enterprise immutable audit log (hash-chained, /security/audit)
        try:
            from security.audit import log_audit as sec_log_audit, AuditCategory, AuditSeverity
            await sec_log_audit(
                action="document_upload",
                category=AuditCategory.DOCUMENT,
                severity=AuditSeverity.INFO,
                actor_id=user_id or owner,
                actor_ip=ip_address or "",
                resource_type="document",
                resource_id=doc.document_id,
                description=f"Uploaded {filename} classified as {doc_category.value}",
            )
        except Exception as e:
            logger.warning("Enterprise audit log write failed (non-fatal): %s", e)

        logger.info(
            "Pipeline COMPLETE — document_id=%s category=%s verified=%s",
            doc.document_id, doc_category.value, verification.value,
        )
        return doc

    # ── Retrieval ────────────────────────────────────────────────────────

    async def retrieve(self, document_id: str) -> bytes:
        """
        Retrieve and decrypt a document from the vault.

        Returns:
            Decrypted file bytes.
        """
        doc = await self.db.get_document(document_id)
        if not doc:
            raise FileNotFoundError(f"Document not found: {document_id}")
        if doc.lifecycle == DocumentLifecycle.DELETED:
            raise PermissionError("Document has been deleted")

        encrypted_blob, wrapped_dek = self.vault.retrieve(
            document_id, doc.encryption_ref
        )
        return self.encryptor.decrypt(encrypted_blob, wrapped_dek)

    # ── Pipeline steps (private) ─────────────────────────────────────────

    def _extract_text(
        self, file_bytes: bytes, mime_type: str, filename: str
    ) -> Tuple[str, float, int]:
        """
        Extract text via OCR (images) or direct text extraction (PDF/DOCX).

        Returns:
            (text, confidence, page_count)
        """
        text = ""
        confidence = 0.0
        page_count = 1

        try:
            if mime_type.startswith("image/"):
                from document.ocr_processor import OCRProcessor
                from PIL import Image

                image = Image.open(io.BytesIO(file_bytes))
                processor = OCRProcessor()
                result = processor.extract_from_image(image, page_number=1)
                text = result["text"]
                confidence = result["confidence"]
                page_count = 1

            elif mime_type == "application/pdf":
                try:
                    import fitz
                    doc = fitz.open(stream=file_bytes, filetype="pdf")
                    page_count = len(doc)
                    pages = []
                    for page in doc:
                        pages.append(page.get_text("text"))
                    text = "\n\n".join(pages)
                    doc.close()
                    confidence = 0.85 if text.strip() else 0.0

                    # If barely any text, try OCR on pages
                    if len(text.strip()) < 100 and page_count > 0:
                        from document.ocr_processor import OCRProcessor
                        from PIL import Image

                        processor = OCRProcessor()
                        doc = fitz.open(stream=file_bytes, filetype="pdf")
                        ocr_texts = []
                        confs = []
                        for pg in doc:
                            pix = pg.get_pixmap(matrix=fitz.Matrix(2, 2))
                            img = Image.open(io.BytesIO(pix.tobytes("png")))
                            r = processor.extract_from_image(img)
                            ocr_texts.append(r["text"])
                            confs.append(r["confidence"])
                        doc.close()
                        text = "\n\n".join(ocr_texts)
                        confidence = sum(confs) / len(confs) if confs else 0.0
                except ImportError:
                    logger.warning("PyMuPDF not available for PDF processing")

            elif "wordprocessingml" in mime_type or mime_type == "application/msword":
                try:
                    import docx as python_docx
                    doc_obj = python_docx.Document(io.BytesIO(file_bytes))
                    text = "\n\n".join(
                        p.text for p in doc_obj.paragraphs if p.text.strip()
                    )
                    confidence = 0.90 if text.strip() else 0.0
                except ImportError:
                    logger.warning("python-docx not available for DOCX processing")

        except Exception as e:
            logger.error("Text extraction failed for '%s': %s", filename, e)

        return text, confidence, page_count

    def _classify(self, text: str) -> Tuple[DocumentCategory, float]:
        """
        Classify OCR text into a DocumentCategory.
        Falls back to keyword matching if the ML classifier isn't available.
        """
        if not text or not text.strip():
            return DocumentCategory.OTHER, 0.0

        # Keyword-based classification (always available)
        text_lower = text.lower()
        scores: Dict[DocumentCategory, float] = {}

        # Load keywords from category schemas
        for cat in DocumentCategory:
            if cat == DocumentCategory.OTHER:
                continue
            schema_path = os.path.join(_CATEGORY_DIR, f"{cat.value}.json")
            if os.path.exists(schema_path):
                with open(schema_path) as f:
                    schema = json.load(f)
                keywords = schema.get("keywords", [])
                hits = sum(1 for kw in keywords if kw.lower() in text_lower)
                if hits > 0:
                    scores[cat] = hits / len(keywords)

        if scores:
            best_cat = max(scores, key=scores.get)
            best_score = scores[best_cat]
            # Normalise to 0-1 range
            confidence = min(best_score * 2, 1.0)
            return best_cat, confidence

        return DocumentCategory.OTHER, 0.1

    def _generate_qr(self, document_id: str) -> str:
        """Generate a QR code containing the document verification URL."""
        try:
            import qrcode

            verification_url = (
                f"https://digilocker.gov.in/verify/{document_id}"
            )
            qr = qrcode.QRCode(version=1, box_size=6, border=2)
            qr.add_data(verification_url)
            qr.make(fit=True)
            img = qr.make_image(fill_color="black", back_color="white")

            buf = io.BytesIO()
            img.save(buf, format="PNG")
            return base64.b64encode(buf.getvalue()).decode()

        except ImportError:
            logger.warning("qrcode library not installed — skipping QR generation")
            return ""
        except Exception as e:
            logger.error("QR generation failed: %s", e)
            return ""

    def _extract_metadata(
        self, text: str, category: DocumentCategory
    ) -> Dict[str, Any]:
        """
        Extract structured metadata fields from OCR text based on category schema.
        Uses regex patterns for common field formats.
        """
        metadata: Dict[str, Any] = {"category_schema": category.value}

        if not text:
            return metadata

        schema_path = os.path.join(_CATEGORY_DIR, f"{category.value}.json")
        if not os.path.exists(schema_path):
            return metadata

        with open(schema_path) as f:
            schema = json.load(f)

        # Basic regex extraction for common patterns
        import re

        text_clean = text.strip()

        # Date patterns
        date_pattern = r'\b(\d{1,2}[/\-\.]\d{1,2}[/\-\.]\d{2,4})\b'
        dates = re.findall(date_pattern, text_clean)
        if dates:
            metadata["extracted_dates"] = dates[:5]

        # Number patterns (could be serial numbers, amounts, etc.)
        number_pattern = r'\b([A-Z]{1,3}[\s\-]?\d{4,12})\b'
        numbers = re.findall(number_pattern, text_clean)
        if numbers:
            metadata["extracted_references"] = numbers[:5]

        # Amount / currency patterns
        amount_pattern = r'(?:Rs\.?|₹|INR)\s*([\d,]+(?:\.\d{2})?)'
        amounts = re.findall(amount_pattern, text_clean)
        if amounts:
            metadata["extracted_amounts"] = amounts[:3]

        # Name-like patterns (after common labels)
        name_pattern = r'(?:Name|name)\s*[:\-]\s*([A-Z][a-zA-Z\s\.]{2,40})'
        names = re.findall(name_pattern, text_clean)
        if names:
            metadata["extracted_names"] = [n.strip() for n in names[:3]]

        return metadata

    # ── Utility ──────────────────────────────────────────────────────────

    @staticmethod
    def _detect_mime(filename: str) -> str:
        """Detect MIME type from filename extension."""
        mime, _ = mimetypes.guess_type(filename)
        if mime:
            return mime
        ext = os.path.splitext(filename)[1].lower()
        ext_map = {
            ".pdf":  "application/pdf",
            ".jpg":  "image/jpeg",
            ".jpeg": "image/jpeg",
            ".png":  "image/png",
            ".docx": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        }
        return ext_map.get(ext, "application/octet-stream")

    @staticmethod
    def get_supported_categories() -> list:
        """Return all supported document categories with their schemas."""
        categories = []
        for cat in DocumentCategory:
            info = {"value": cat.value, "display_name": cat.value.replace("_", " ").title()}
            schema_path = os.path.join(_CATEGORY_DIR, f"{cat.value}.json")
            if os.path.exists(schema_path):
                with open(schema_path) as f:
                    schema = json.load(f)
                info["display_name"] = schema.get("display_name", info["display_name"])
                info["department"] = schema.get("department", "")
                info["description"] = schema.get("description", "")
                info["field_count"] = len(schema.get("fields", []))
            categories.append(info)
        return categories
